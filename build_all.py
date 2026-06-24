#!/usr/bin/env python3
"""
build_all.py -- single reproducible build for the thesis.

Guarantees nothing stale reaches any output: PDF, docx (16 components + combined),
and the LaTeX bundle are ALL regenerated from the canonical .md / .bib every run.
Tables are rebuilt from scratch in python-docx (pandoc's docx tables collapse), and
a render-and-check assertion verifies every table's column count survives into the
rendered docx -- so a silent table collapse can never reach the reviewer again.

Run:  python3 build_all.py
"""
import os, re, sys, glob, shutil, subprocess, json
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

OUT   = os.path.dirname(os.path.abspath(__file__))
BUILD = os.path.join(OUT, "mitthesis_build")
DEL   = os.path.join(OUT, "thesis_deliverables")
DOCX  = os.path.join(DEL, "docx")
MD    = os.path.join(DEL, "md")
TEX   = os.path.join(DEL, "tex")
REF   = os.path.join(DEL, "_reference.docx")

# canonical source -> deliverable component  (cover + references handled separately)
def _latest(prefix):
    """Newest _vN.md for a stem. Same rule build.py uses, so the docx/PDF/tex builds all
    read the identical newest version and the two source maps cannot drift apart."""
    cands = glob.glob(os.path.join(OUT, prefix + "_v*.md"))
    if not cands:
        raise RuntimeError("no source matching %s_v*.md in %s" % (prefix, OUT))
    return os.path.basename(max(cands, key=lambda p: int(re.search(r"_v(\d+)\.md$", p).group(1))))

MAP = [
 ("01_Abstract", _latest("Abstract_Ekekwe")),
 ("02_Chapter_1_Introduction", _latest("Chapter1_Full")),
 ("03_Chapter_2_Theoretical_Foundation", _latest("Chapter2_Full")),
 ("04_Chapter_3_Research_Methodology", _latest("Chapter3_Full")),
 ("05_Chapter_4_Empirical_Analysis", _latest("Chapter4_Full")),
 ("06_Chapter_5_Capability_Assurance_Framework", _latest("Chapter5_Full")),
 ("07_Chapter_6_Conclusions", _latest("Chapter6_Full")),
 ("08_Appendix_A_Case_Study_Protocol", "Appendix_A_Case_Study_Protocol.md"),
 ("09_Appendix_B_Discrimination_Matrix", "Appendix_B_Discrimination_Matrix.md"),
 ("10_Appendix_C_Interview_Guides", "Appendix_C_Interview_Guides.md"),
 ("11_Appendix_D_Case_by_Case_Evidence", "Appendix_D_Case_by_Case_Evidence_Table.md"),
 ("12_Appendix_E_Verdict_Traceability_Matrix", "Appendix_E_Verdict_Traceability_Matrix.md"),
 ("13_Appendix_F_Survey_Response_Data", "Appendix_F_Survey_Response_Data.md"),
 ("14_Appendix_G_Capability_Element_Traceability_Matrix", "Appendix_G_Capability_Element_Traceability_Matrix.md"),
]
ORDER = ["00_Cover_Page"] + [m[0] for m in MAP] + ["15_References"]

# ============================ number-drift guard ============================
# Standing PRE-BUILD check. Extract every digit-bearing figure from each source,
# diff the DOCUMENT-WIDE set against the previous successful build, and flag any
# figure that has dropped out of the WHOLE document. Design choices:
#   * Document-wide (all 14 MAP sources, not per chapter) so a figure relocated
#     chapter->appendix or chapter->chapter is NOT a false alarm -- only a true
#     disappearance flags. This is the case that motivated the guard.
#   * The unit stays attached to the number ("30 years" is its own token, distinct
#     from the "30" in "OML 30"), so a unit-bearing figure cannot be masked by a
#     bare integer of the same magnitude living elsewhere.
#   * NON-BLOCKING. A dropped figure is often an intentional trim; the guard warns
#     and lets the build proceed. The snapshot is committed only after a clean build.
#   * Scope is digit-bearing tokens (currency, percent, scaled magnitudes, numbers
#     with units, bare numbers/years). Spelled-out numbers ("seven workers") are NOT
#     tracked -- matching them floods the set with ordinary prose words.
SNAPSHOT = os.path.join(OUT, "_number_snapshot.json")
_SCALE = r"million|billion|trillion|thousand"
_UNIT  = r"barrels|bpd|man-hours|man\s?hours|tonnes?|days?|years?|hours?|months?|weeks?|fields?|communities|wells?"
_NUMTOK = re.compile(
    r"(?P<money>[$\u00a3\u20ac]\s?\d[\d,]*(?:\.\d+)?(?:\s+(?:" + _SCALE + r"))?)"
    r"|(?P<pct>\d[\d,]*(?:\.\d+)?\s+percent)"
    r"|(?P<scaled>\d[\d,]*(?:\.\d+)?\s+(?:" + _SCALE + r"))"
    r"|(?P<united>\d[\d,]*(?:\.\d+)?\s+(?:" + _UNIT + r"))"
    r"|(?P<bare>\d[\d,]*(?:\.\d+)?)", re.IGNORECASE)

def _norm_num(tok):
    t = re.sub(r"\s+", " ", tok.strip().lower()).replace(",", "")
    return re.sub(r"([$\u00a3\u20ac])\s+", r"\1", t).replace("man hours", "man-hours")

def _strip_md(s):
    return "\n".join(ln for ln in s.split("\n")
                     if not ln.strip().startswith(("![", "<!--", "```")))

def extract_numbers(text):
    flat = re.sub(r"\s+", " ", text); found = {}
    for m in _NUMTOK.finditer(flat):
        n = _norm_num(m.group(0))
        if n not in found:
            a = max(0, m.start() - 55); b = min(len(flat), m.end() + 55)
            found[n] = "..." + flat[a:b].strip() + "..."
    return found

def _label(stem):
    m = re.search(r"Chapter_(\d)", stem)
    if m: return "Ch" + m.group(1)
    m = re.search(r"Appendix_([A-Z])", stem)
    if m: return "Appendix " + m.group(1)
    return "Abstract" if "Abstract" in stem else stem

def collect_numbers():
    doc = {}
    for stem, src in MAP:
        p = os.path.join(OUT, src)
        if not os.path.exists(p): continue
        lab = _label(stem)
        for n, ctx in extract_numbers(_strip_md(open(p, encoding="utf-8").read())).items():
            doc.setdefault(n, {}).setdefault(lab, ctx)
    return doc

def _classify(t):
    if re.match(r"^[$\u00a3\u20ac]", t): return "money"
    if t.endswith("percent"): return "pct"
    if re.search(r"(million|billion|trillion|thousand)$", t): return "scaled"
    if re.search(r"(barrels|bpd|man-hours|tonnes?|days?|years?|hours?|months?|weeks?|fields?|communities|wells?)$", t): return "unit"
    if re.match(r"^(19|20)\d\d$", t): return "year"
    return "bare"

def check_number_drift():
    """Pre-build, non-blocking. Returns the current figure set to snapshot on success."""
    print("[0/5] NUMBERS (figures vs previous build: catch any figure vanishing from the whole document)")
    current = collect_numbers()
    prev = None
    if os.path.exists(SNAPSHOT):
        try: prev = json.load(open(SNAPSHOT, encoding="utf-8")).get("tokens")
        except Exception as e: print("     (snapshot unreadable: %s -- re-establishing baseline)" % e)
    if not prev:
        nf = len({l for v in current.values() for l in v})
        print("     baseline established: %d distinct figures across %d files (no previous build to diff)" % (len(current), nf))
        return current
    dropped = [t for t in prev if t not in current]
    added   = [t for t in current if t not in prev]
    print("     current %d figures; previous %d" % (len(current), len(prev)))
    if not dropped:
        print("     OK: no figure dropped from the document since last build" + ((" (%d added)" % len(added)) if added else ""))
        return current
    rank = {"money":0,"pct":1,"scaled":2,"unit":3,"bare":4,"year":5}
    dropped.sort(key=lambda t: (rank.get(_classify(t), 9), t))
    print("     >> %d figure(s) DROPPED ENTIRELY since last build -- confirm each is intentional:" % len(dropped))
    for t in dropped:
        where = prev[t]; lab = next(iter(where))
        print("        [%-5s] %-22s was in %s: %s" % (_classify(t), t, lab, where[lab][:88]))
    if added:
        print("     (%d figure token(s) added since last build; additions are not loss, not listed)" % len(added))
    print("     NOTE non-blocking. A small bare integer can collide with a section/table/OML number; read the context.")
    return current

def save_number_snapshot(current):
    import datetime
    json.dump({"_meta": {"timestamp": datetime.datetime.now().isoformat(timespec="seconds"),
                          "token_count": len(current)}, "tokens": current},
              open(SNAPSHOT, "w", encoding="utf-8"), indent=1, ensure_ascii=False)
    print("     number snapshot updated (%d figures) -> %s" % (len(current), os.path.basename(SNAPSHOT)))


def page_twips():
    s = Document(REF).sections[0]
    return int((int(s.page_width) - int(s.left_margin) - int(s.right_margin)) / 635)

# ---------- table rebuild (pandoc tables are unrenderable; build clean ones) ----------
def _edge(tcB, name, val, sz=None):
    el = OxmlElement(f'w:{name}')
    if val == 'nil':
        el.set(qn('w:val'), 'nil')
    else:
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0');    el.set(qn('w:color'), '000000')
    tcB.append(el)

def _word_w(lw, fs):
    """Twips an unbreakable word of lw chars needs at fs pt, including cell padding."""
    return int(lw * fs * 10) + 200

def _longest_word(grid, c, nr):
    return max((max((len(w) for w in grid[r][c].split()), default=1) for r in range(nr)), default=1)

def _col_min(grid, c, nr, fs):
    return max(_word_w(_longest_word(grid, c, nr), fs), 520)

# ---------- landscape wrapping for wide tables (match the PDF) ----------
def _body_dims(doc):
    sect = doc.element.body.find(qn('w:sectPr'))
    sz = sect.find(qn('w:pgSz')); mar = sect.find(qn('w:pgMar'))
    g = lambda e, k: int(e.get(qn(f'w:{k}')))
    return (g(sz, 'w'), g(sz, 'h'), g(mar, 'top'), g(mar, 'bottom'), g(mar, 'left'), g(mar, 'right'))

def _make_sectpr(dims, landscape):
    pw, ph, top, bot, left, right = dims
    s = OxmlElement('w:sectPr'); sz = OxmlElement('w:pgSz')
    if landscape:
        sz.set(qn('w:w'), str(ph)); sz.set(qn('w:h'), str(pw)); sz.set(qn('w:orient'), 'landscape')
    else:
        sz.set(qn('w:w'), str(pw)); sz.set(qn('w:h'), str(ph))
    s.append(sz); m = OxmlElement('w:pgMar')
    for k, v in [('top', top), ('bottom', bot), ('left', left), ('right', right),
                 ('header', 720), ('footer', 720), ('gutter', 0)]:
        m.set(qn(f'w:{k}'), str(v))
    s.append(m); return s

def _wrap_landscape(doc, tbl, dims):
    """Put tbl (and its caption) in a landscape section between two portrait sections."""
    cap = tbl.getprevious()                                   # caption paragraph
    anchor = cap.getprevious() if (cap is not None and cap.tag == qn('w:p')) else cap
    if anchor is None or anchor.tag != qn('w:p'):
        anchor = OxmlElement('w:p'); (cap if cap is not None else tbl).addprevious(anchor)
    pPr = anchor.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr'); anchor.insert(0, pPr)
    if pPr.find(qn('w:sectPr')) is None:
        pPr.append(_make_sectpr(dims, False))                 # close portrait section
    nxt = OxmlElement('w:p'); npr = OxmlElement('w:pPr')
    npr.append(_make_sectpr(dims, True)); nxt.append(npr)     # close landscape section after table
    tbl.addnext(nxt)

def build_clean_table(doc, grid, total, fs):
    nr, nc = len(grid), len(grid[0])
    mins = [_col_min(grid, c, nr, fs) for c in range(nc)]
    wts  = [min(max(max(len(grid[r][c]) for r in range(nr)), 3), 80) for c in range(nc)]
    rem  = total - sum(mins)
    if rem > 0:
        widths = [mins[c] + int(rem * wts[c] / sum(wts)) for c in range(nc)]
    else:
        sc = total / sum(mins); widths = [int(mins[c] * sc) for c in range(nc)]
    widths[-1] = total - sum(widths[:-1])
    tbl = doc.add_table(rows=nr, cols=nc)
    tblPr = tbl._tbl.tblPr
    lay = OxmlElement('w:tblLayout'); lay.set(qn('w:type'), 'fixed'); tblPr.append(lay)
    tW  = OxmlElement('w:tblW'); tW.set(qn('w:type'), 'dxa'); tW.set(qn('w:w'), str(total)); tblPr.append(tW)
    g = tbl._tbl.find(qn('w:tblGrid'))
    for gc in list(g):
        g.remove(gc)
    for w in widths:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(w)); g.append(gc)
    for r in range(nr):
        for c in range(nc):
            cell = tbl.cell(r, c); cell.text = grid[r][c]
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW'); tcW.set(qn('w:type'), 'dxa'); tcW.set(qn('w:w'), str(widths[c])); tcPr.append(tcW)
            tcB = OxmlElement('w:tcBorders'); tcPr.append(tcB)
            _edge(tcB, 'left', 'nil'); _edge(tcB, 'right', 'nil')
            _edge(tcB, 'top', 'single', 8) if r == 0 else _edge(tcB, 'top', 'nil')
            if r == nr - 1:   _edge(tcB, 'bottom', 'single', 8)
            elif r == 0:      _edge(tcB, 'bottom', 'single', 4)
            else:             _edge(tcB, 'bottom', 'nil')
            for p in cell.paragraphs:
                pf = p.paragraph_format; pf.space_before = Pt(2); pf.space_after = Pt(2); pf.line_spacing = 1.0
                for run in p.runs:
                    run.font.size = Pt(fs)
                    if r == 0:
                        run.font.bold = True
    return tbl

def _layout(grid, port_w, land_w):
    """Decide (font, total_width, landscape) so no column is narrower than its longest word."""
    nr, nc = len(grid), len(grid[0])
    if nc < 6:                                   # try portrait at 9 then 8 pt
        for fs in (9, 8):
            if sum(_col_min(grid, c, nr, fs) for c in range(nc)) <= port_w:
                return fs, port_w, False
    for fs in (8, 7):                            # widest / non-fitting tables go landscape
        if sum(_col_min(grid, c, nr, fs) for c in range(nc)) <= land_w:
            return fs, land_w, True
    return 7, land_w, True

def rebuild_tables(path, port_w):
    doc = Document(path)
    dims = _body_dims(doc)
    land_w = dims[1] - dims[4] - dims[5]            # landscape text width = pageH - left - right
    olds  = list(doc.tables)
    grids = [[[c.text for c in row.cells] for row in t.rows] for t in olds]
    for idx, old in enumerate(olds):
        oe = old._tbl
        fs, total, landscape = _layout(grids[idx], port_w, land_w)
        newt = build_clean_table(doc, grids[idx], total, fs)
        oe.addprevious(newt._tbl)
        if landscape:
            _wrap_landscape(doc, newt._tbl, dims)
        oe.getparent().remove(oe)
    doc.save(path)
    return len(olds)

# ---------- markdown pipe-table parser (for the assertion) ----------
def parse_md_tables(md):
    """Return [[header cells], ...] for each pipe table in source order."""
    out, lines, i = [], md.split("\n"), 0
    sep = re.compile(r'^\s*\|?[\s:|-]*-[\s:|-]*\|?\s*$')
    while i < len(lines):
        ln = lines[i]
        if ln.lstrip().startswith("|") and i + 1 < len(lines) and "---" in lines[i + 1] and sep.match(lines[i + 1]):
            hdr = [c.strip() for c in ln.strip().strip("|").split("|")]
            out.append(hdr)
            i += 2
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                i += 1
        else:
            i += 1
    return out

# ============================ build steps ============================
def build_pdf():
    print("[1/5] PDF  (chapters, appendices, figures, bib, abstract -> all synced from canonical)")
    r = subprocess.run([sys.executable, "build.py"], cwd=BUILD, capture_output=True, text=True)
    for line in r.stdout.splitlines():
        if any(k in line for k in ("synced", "compiled", "regenerated", "WARNING", "check:")):
            print("     " + line.strip())
    if r.returncode != 0:
        sys.stderr.write(r.stdout + "\n" + r.stderr + "\n")
        raise RuntimeError("build.py failed (a float/numbering/reference guard tripped) -- "
                           "PDF NOT updated; see output above")
    shutil.copy(os.path.join(BUILD, "MIT_format_proof.pdf"), os.path.join(DEL, "Thesis_Full_Proof.pdf"))

def build_docx(total):
    print("[2/5] docx (16 components, regenerated from canonical, tables rebuilt from scratch)")
    os.chdir(OUT)
    ntab = 0
    for stem, src in MAP:
        md = open(src, encoding="utf-8").read().replace(".svg", ".png")   # PNG embeds reliably
        tmp = os.path.join("/tmp", stem + ".md"); open(tmp, "w", encoding="utf-8").write(md)
        open(os.path.join(MD, stem + ".md"), "w", encoding="utf-8").write(md)
        out = os.path.join(DOCX, stem + ".docx")
        subprocess.run(["pandoc", tmp, "--reference-doc", REF, "-o", out], check=True,
                       capture_output=True, text=True)
        ntab += rebuild_tables(out, total)
    # References from normalized bib
    b = open(os.path.join(OUT, "references.bib")).read()
    b = re.sub(r'journaltitle\s*=', 'journal =', b)
    b = re.sub(r'\bdate\s*=\s*\{(\d{4})[^}]*\}', r'year = {\1}', b)
    nb = os.path.join(DEL, "_refn.bib"); open(nb, "w").write(b)
    st = os.path.join(DEL, "_refs.md"); open(st, "w").write("---\nnocite: '@*'\n---\n\n# References\n")
    subprocess.run(["pandoc", st, "--citeproc", "--bibliography", nb, "--reference-doc", REF,
                    "-o", os.path.join(DOCX, "15_References.docx")], check=True)
    subprocess.run(["pandoc", st, "--citeproc", "--bibliography", nb, "-t", "markdown",
                    "-o", os.path.join(MD, "15_References.md")], check=True)
    # combined md + docx
    parts = [open(os.path.join(MD, x + ".md")).read().strip() for x in ORDER
             if os.path.exists(os.path.join(MD, x + ".md"))]
    combined = "\n\n\\newpage\n\n".join(parts).replace(".svg", ".png")
    open(os.path.join(DEL, "Thesis_Full_Continuous.md"), "w").write(combined)
    ct = "/tmp/_combined.md"; open(ct, "w").write(combined)
    subprocess.run(["pandoc", ct, "--reference-doc", REF, "-o", os.path.join(DEL, "Thesis_Full.docx")], check=True)
    ntab += rebuild_tables(os.path.join(DEL, "Thesis_Full.docx"), total)
    for f in ("_refn.bib", "_refs.md"):
        p = os.path.join(DEL, f); os.path.exists(p) and os.remove(p)
    print(f"     {ntab} tables rebuilt from scratch (components + combined)")

def build_tex():
    print("[3/5] tex  (modular build source + continuous single file, from canonical build)")
    os.makedirs(TEX, exist_ok=True)
    for old in glob.glob(os.path.join(TEX, "*")):
        os.remove(old)
    for pat in ("*.tex", "*.cls", "*.bib", "*.sty"):
        for f in glob.glob(os.path.join(BUILD, pat)):
            shutil.copy(f, TEX)
    for png in glob.glob(os.path.join(OUT, "Figure_*.png")):
        shutil.copy(png, TEX)
    # continuous single-file .tex: inline \include / \input
    main = open(os.path.join(BUILD, "main.tex")).read()
    def inline(m):
        p = os.path.join(BUILD, m.group(1) + ".tex")
        return ("\n% --- " + m.group(1) + " ---\n" + open(p).read()) if os.path.exists(p) else m.group(0)
    main = re.sub(r'\\(?:include|input)\{([^}]+)\}', inline, main)
    open(os.path.join(TEX, "Thesis_Full_Continuous.tex"), "w").write(main)
    shutil.copy(os.path.join(BUILD, "MIT_format_proof.pdf"), os.path.join(TEX, "MIT_format_proof.pdf"))

def _clean_tokens(hdr):
    """Distinctive header tokens (strip markdown, first word, len>=4) for render matching."""
    toks = []
    for h in hdr:
        w = re.sub(r'[*`|]', '', h).strip()
        if w and len(w.split()[0]) >= 4:
            toks.append(w.split()[0])
    return toks

def check_tables(total):
    print("[4/5] CHECK (md vs docx: count + columns + fixed-layout; render: no collapse)")
    ncomp = 0
    # (a) per-component XML check: table count, column count, and fixed layout (the collapse cause)
    for stem, src in MAP:
        md_tabs = parse_md_tables(open(src).read())
        d = Document(os.path.join(DOCX, stem + ".docx"))
        dx = [len(t.columns) for t in d.tables]
        assert len(md_tabs) == len(dx), f"{stem}: table COUNT md={len(md_tabs)} docx={len(dx)}"
        for i, (mc, dc) in enumerate(zip([len(h) for h in md_tabs], dx)):
            assert mc == dc, f"{stem} table {i}: COLUMN md={mc} docx={dc}"
        for i, t in enumerate(d.tables):
            lay = t._tbl.tblPr.find(qn('w:tblLayout'))
            v = lay.get(qn('w:type')) if lay is not None else 'none'
            assert v == 'fixed', f"{stem} table {i}: layout='{v}' not 'fixed' (pandoc table -> would collapse)"
            # column-width sanity: no column narrower than its longest unbreakable word
            grid = [[c.text for c in row.cells] for row in t.rows]
            fs = 9
            for p in t.cell(0, 0).paragraphs:
                rs = [r for r in p.runs if r.font.size]
                if rs:
                    fs = int(round(rs[0].font.size.pt)); break
            widths = [int(gc.get(qn('w:w'))) for gc in t._tbl.find(qn('w:tblGrid')).findall(qn('w:gridCol'))]
            for c in range(len(grid[0])):
                lw = _longest_word(grid, c, len(grid)); need = _word_w(lw, fs)
                assert widths[c] + 40 >= need, (
                    f"{stem} table {i} col {c}: width {widths[c]} < longest-word need {need} "
                    f"({lw}-char word at {fs}pt would wrap)")
        ncomp += len(md_tabs)
    # (b) render-check the combined docx: every table must render as a grid (not a stacked column)
    rdir = "/tmp/_rendercheck"; shutil.rmtree(rdir, ignore_errors=True); os.makedirs(rdir)
    full = os.path.join(DEL, "Thesis_Full.docx")
    subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", rdir, full],
                   capture_output=True, timeout=300)
    layout = subprocess.run(["pdftotext", "-layout", os.path.join(rdir, "Thesis_Full.pdf"), "-"],
                            capture_output=True, text=True).stdout
    lines = layout.split("\n")
    md_tabs = parse_md_tables(open(os.path.join(DEL, "Thesis_Full_Continuous.md")).read())
    collapsed = []
    for ti, hdr in enumerate(md_tabs):
        toks = _clean_tokens(hdr)
        if len(toks) < 2:
            continue
        # a grid puts >=2 header tokens on one rendered line; a collapse stacks each cell alone
        if not any(sum(1 for t in toks if t in ln) >= 2 for ln in lines):
            collapsed.append((ti, [re.sub(r'[*`]', '', h) for h in hdr[:3]]))
    assert not collapsed, f"COLLAPSED tables in render: {collapsed}"
    print(f"     OK: {ncomp} component tables fixed-layout + column-matched; "
          f"{len(md_tabs)} combined tables render as grids")

def make_zip():
    print("[5/5] zip")
    os.chdir(DEL)
    z = os.path.join(OUT, "thesis_deliverables.zip")
    if os.path.exists(z):
        os.remove(z)
    subprocess.run(["zip", "-qr", z, ".", "-x", "_*", "*.DS_Store"], capture_output=True)
    print("     thesis_deliverables.zip written")

if __name__ == "__main__":
    print("=== reproducible canonical build: PDF + docx + tex, self-checked ===")
    nums = check_number_drift()          # pre-build, non-blocking figure-drift warning
    tw = page_twips()
    build_pdf()
    build_docx(tw)
    build_tex()
    check_tables(tw)
    make_zip()
    save_number_snapshot(nums)           # commit snapshot only after a clean build
    print("=== done: all outputs sourced from canonical; table render-check passed ===")
